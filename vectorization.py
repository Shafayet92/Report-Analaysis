import hashlib
import os
import re
from sentence_transformers import CrossEncoder
import torch
import pandas as pd
import logging
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain.schema import Document
import numpy as np
from docx import Document as DocxDocument  # For handling DOCX files
from typing import List, Tuple, Any

from search import fullsummarization, llm_response_search, summarize_data

UPLOAD_FOLDER = './uploads'
PERSIST_DIRECTORY = 'data/db'
SIMILARITY_THRESHOLD = 0.3  # Cosine similarity threshold
BATCH_SIZE = 1000  # Batch size for processing

# Initialize the reranker (outside the class for reuse)
reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')

# Ensure folders exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PERSIST_DIRECTORY, exist_ok=True)

logging.basicConfig(level=logging.INFO)


def load_pdf(file_path):
    """ Load PDF and add file name metadata """
    loader = PyPDFLoader(file_path)
    documents = loader.load()
    file_name = os.path.basename(file_path)

    # Attach metadata to every document chunk
    for doc in documents:
        doc.metadata = {"file_name": file_name}  # Ensure metadata is stored

    return documents


# Function to load Excel files
def load_excel(file_path):
    """ Load Excel and add file name metadata """
    df = pd.read_excel(file_path)
    documents = []
    file_name = os.path.basename(file_path)

    for i in range(0, len(df), BATCH_SIZE):
        chunk = df.iloc[i:i + BATCH_SIZE]
        text = chunk.to_string(index=False)
        documents.append(Document(page_content=text, metadata={"file_name": file_name}))

    return documents


def load_csv(file_path):
    """ Load CSV and add file name metadata """
    df = pd.read_csv(file_path)
    documents = []
    file_name = os.path.basename(file_path)

    for i in range(0, len(df), BATCH_SIZE):
        chunk = df.iloc[i:i + BATCH_SIZE]
        text = chunk.to_string(index=False)
        documents.append(Document(page_content=text, metadata={"file_name": file_name}))

    return documents

# Function to load DOCX files
def load_docx(file_path):
    """ Load DOCX and add file name metadata """
    doc = DocxDocument(file_path)
    file_name = os.path.basename(file_path)

    documents = [Document(page_content=para.text, metadata={"file_name": file_name})
                 for para in doc.paragraphs if para.text.strip()]
    return documents

# Function to split documents into smaller chunks
# def text_split(extracted_data, chunk_size=1000, chunk_overlap=150):

#     separators = ["\n\n", "\n", ". ", " ", ""]

#     text_splitter = RecursiveCharacterTextSplitter(
#         chunk_size=chunk_size,
#         chunk_overlap=chunk_overlap,
#         separators=separators
#     )

#     # Split the document
#     text_chunks = text_splitter.split_documents(extracted_data)

#     # Remove empty or redundant chunks
#     filtered_chunks = [chunk for chunk in text_chunks if len(chunk.page_content.strip()) > 0]

#     return filtered_chunks

def text_split(extracted_data, max_sentences=3, chunk_size=1000, overlap=1):
    """
    Hybrid chunking: Sentence-based but ensures chunk size is reasonable.
    :param extracted_data: List of Document objects.
    :param max_sentences: Number of sentences per chunk.
    :param chunk_size: Character limit for each chunk.
    :param overlap: Number of overlapping sentences.
    """
    all_chunks = []

    for doc in extracted_data:
        text = doc.page_content.strip()
        file_name = doc.metadata.get("file_name", "Unknown")

        # Step 1: Split into full sentences
        sentences = re.split(r'(?<=[.!?])\s+', text)

        # Step 2: Combine sentences into chunks, ensuring they don't exceed chunk_size
        current_chunk = []
        current_length = 0

        for sentence in sentences:
            if current_length + len(sentence) > chunk_size and current_chunk:
                chunk_text = " ".join(current_chunk)
                all_chunks.append(Document(page_content=chunk_text, metadata={"file_name": file_name}))
                current_chunk = current_chunk[-overlap:]  # Retain overlap
                current_length = sum(len(s) for s in current_chunk)

            current_chunk.append(sentence)
            current_length += len(sentence)

        # Add last chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            all_chunks.append(Document(page_content=chunk_text, metadata={"file_name": file_name}))

    return all_chunks



# Get embedding model
def get_embeddings():
    return HuggingFaceEmbeddings(
        model_name="multi-qa-MiniLM-L6-cos-v1",
        model_kwargs={'device': 'cuda' if torch.cuda.is_available() else 'cpu'}
    )

def cosine_similarity(vec1, vec2):
        vec1 = np.array(vec1)
        vec2 = np.array(vec2)
        norm1 = np.linalg.norm(vec1)
        norm2 = np.linalg.norm(vec2)
        if norm1 == 0 or norm2 == 0:
            return 0.0
        return np.dot(vec1, vec2) / (norm1 * norm2)

# VectorStore class to manage vectorized data
class VectorStore:
    def __init__(self):
        self.embeddings = get_embeddings()
        self.vectordb = None
        self.document_hashes = set()  # Set to store hashes of existing documents
        self.initialize_store()

    def initialize_store(self):
        # Initialize the Chroma vector store with existing data if available
        if os.path.exists(PERSIST_DIRECTORY):
            self.vectordb = Chroma(
                persist_directory=PERSIST_DIRECTORY,
                embedding_function=self.embeddings
            )
            # Load existing document hashes from the vector store
            all_docs = self.vectordb.get()
            for doc in all_docs:
                # Ensure doc is a Document object before trying to access page_content
                if isinstance(doc, Document):
                    doc_hash = hashlib.md5(doc.page_content.encode('utf-8')).hexdigest()
                    self.document_hashes.add(doc_hash)
                else:
                    logging.warning(f"Skipping non-Document object: {doc}")
        else:
            self.vectordb = Chroma(
                embedding_function=self.embeddings,
                persist_directory=PERSIST_DIRECTORY
            )

    def clear_store(self):
        """ Clears the vector store data, useful on a system restart if you want a fresh start """
        self.vectordb.clear()
        print("Cleared the Chroma vector store.")  # Debugging log

    def add_documents(self, documents):
        # Process documents in batches
        text_chunks = text_split(documents)
        total_chunks = len(text_chunks)

        # Add documents to the vector store in batches
        for i in range(0, total_chunks, BATCH_SIZE):
            batch = text_chunks[i:i + BATCH_SIZE]
            unique_batch = []

            for doc in batch:
                doc_hash = hashlib.md5(doc.page_content.encode('utf-8')).hexdigest()

                # Check if the hash is already in the set of document hashes
                if doc_hash not in self.document_hashes:
                    unique_batch.append(doc)
                    self.document_hashes.add(doc_hash)

            if unique_batch:
                self.vectordb.add_documents(unique_batch)
                print(f"Processed batch {i // BATCH_SIZE + 1}/{(total_chunks + BATCH_SIZE - 1) // BATCH_SIZE}")
            else:
                print(f"Batch {i // BATCH_SIZE + 1} contains only duplicates. Skipping.")

        return total_chunks



    def normalize_score(self, score):
        # Convert cosine similarity to [0, 1] range
        return (score + 1) / 2


    def delete_chunks_by_file(self, file_name: str) -> bool:
        """
        Deletes all document chunks from the vector store that have a metadata field 'file_name'
        matching the provided file_name.

        Args:
            file_name (str): The file name whose chunks should be removed.

        Returns:
            bool: True if deletion was successful, False otherwise.
        """
        try:
            # Use where_document to target documents whose metadata has file_name equal to the provided value.
            self.vectordb.delete(where_document={"file_name": file_name})
            logging.info(f"Deleted all chunks with file_name '{file_name}' from the vector store.")

            # Rebuild the document_hashes set to reflect the current state of the vector store.
            self.document_hashes = set()
            all_docs = self.vectordb.get()
            for doc in all_docs:
                if isinstance(doc, Document):
                    doc_hash = hashlib.md5(doc.page_content.encode('utf-8')).hexdigest()
                    self.document_hashes.add(doc_hash)
            return True

        except Exception as e:
            logging.error(f"Error deleting chunks for file '{file_name}': {e}")
            return False



    def pure_chroma_mode(self, query: str, k: int = 30) -> List[Tuple[Any, str]]:
        """
        Perform a similarity search on the vector database and optionally rerank the results.
        """
        if not self.vectordb:
            logging.warning("Vector store is not initialized. Returning empty results.")
            return []

        # Normalize the query
        normalized_query = query.lower().strip()

        try:
            # Retrieve initial results from the vector store
            raw_results = self.vectordb.similarity_search_with_relevance_scores(normalized_query, k=k)
            if not raw_results:
                logging.info(f"No results returned from the vector store for query: '{normalized_query}'")
                return []

            # Process raw results into valid (doc, score) tuples
            results = [(res[0], res[1]) for res in raw_results if isinstance(res, (tuple, list)) and len(res) >= 2]
            if not results:
                logging.info(f"No valid results found for query: '{normalized_query}'")
                return []

            # Normalize scores to [0, 1] and assign relevancy levels
            ranked_results = []
            for doc, score in results:
                norm_score = (score + 1) / 2  # Normalize from [-1, 1] to [0, 1]

                # Assign relevancy levels based on normalized score
                if norm_score >= 0.75:
                    relevancy = "High"
                elif norm_score >= 0.6:
                    relevancy = "Medium"
                else:
                    relevancy = "Low"

                ranked_results.append((doc, relevancy))

            # Extract page content for reranking
            docs_texts = [getattr(doc, 'page_content', None) for doc, _ in ranked_results if hasattr(doc, 'page_content')]
            if not docs_texts:
                logging.debug("No documents with 'page_content' found for reranking.")

            if hasattr(self, "reranker") and self.reranker and docs_texts:
                # Rerank the documents if reranker is available
                query_doc_pairs = [(normalized_query, doc_text) for doc_text in docs_texts]
                rerank_scores = self.reranker.predict(query_doc_pairs)

                if len(rerank_scores) == len(ranked_results):
                    reranked_results = sorted(
                        [(doc, score) for ((doc, _), score) in zip(ranked_results, rerank_scores)],
                        key=lambda x: x[1], reverse=True
                    )
                    return reranked_results
                else:
                    logging.warning("Mismatch between rerank scores and retrieved documents.")
                    return ranked_results

            # If no reranker, return normalized results
            logging.info("Reranker not available; returning normalized results.")
            return ranked_results

        except Exception as e:
            logging.error(f"Error during similarity search: {e}", exc_info=True)
            return []




    def chroma_and_LLM_mode(self, query: str, initial_k: int = 10, step: int = 10, max_k: int = 150, llm=None) -> List[Tuple[Any, float]]:
        """
        Retrieve results using pure_chroma_mode() and incrementally expand the search using LLM checks.

        This function:
        - Retrieves results using pure_chroma_mode() with an initial `k`.
        - Expands the search in increments of `step` until `max_k` or the LLM deems further results irrelevant.
        - Returns the same format as pure_chroma_mode(), with LLM only used for checking relevance.

        Args:
            query (str): The search query.
            initial_k (int): Initial number of results to retrieve.
            step (int): Increment for each expansion step.
            max_k (int): Maximum number of results to retrieve.
            llm: An LLM instance used for determining whether to expand the search.

        Returns:
            List[Tuple[Any, float]]: A list of tuples, each containing a document and its corresponding score.
        """

        # Step 1: Initial search using pure_chroma_mode() with `initial_k` results
        current_k = initial_k
        results = self.pure_chroma_mode(query, k=current_k)

        # If no results are returned, log the information and return an empty list
        if not results:
            logging.info("No results returned from the initial search.")
            return []

        # Step 2: Start expanding the search by increasing `k` in steps until we reach `max_k`
        while current_k < max_k:

            result_to_check = results[-1]  # Check the first from new or last from old batch

            try:
                # Get the LLM response for checking the relevance of the result to check
                llm_response = llm_response_search(query, [result_to_check], model="llama3.2").strip().lower()
                logging.info(f"LLM Response: {llm_response}")

                # If the LLM response is "no", stop further expansion as additional results are irrelevant
                if llm_response.startswith("no"):
                    logging.info("LLM judged additional result as irrelevant. Stopping expansion.")
                    break
            except Exception as e:
                # If an error occurs during LLM evaluation, log it and break out of the loop
                logging.error(f"Error during LLM evaluation: {e}")
                break

            # Calculate the next value of k, which is the minimum of current_k + step or max_k
            next_k = min(current_k + step, max_k)

            # Step 3: Retrieve additional results using pure_chroma_mode()
            new_results = self.pure_chroma_mode(query, k=next_k)

            # Ensure there are additional results beyond the current_k (no duplication of old results)
            if len(new_results) <= current_k:
                logging.info("No more new results available. Stopping expansion.")
                break

            # Step 4: Check the first result of the next batch or last result of the current batch
            additional_results = new_results[current_k:next_k]  # Get the new results starting from `current_k`

            # Step 5: If LLM confirms the relevance of additional results, expand the results
            results = new_results  # Update results with the newly expanded results
            current_k = next_k  # Update `current_k` to the next batch for the next iteration

        # Return the final results after considering LLM decisions for relevance
        return results



def convert_to_serializable(obj):
    if isinstance(obj, dict):
        return {key: convert_to_serializable(value) for key, value in obj.items()}
    elif isinstance(obj, list):
        return [convert_to_serializable(item) for item in obj]
    elif isinstance(obj, np.float32):  # Specifically handle NumPy float32
        return float(obj)
    elif isinstance(obj, np.ndarray):  # Handle NumPy arrays
        return obj.tolist()
    else:
        return obj


# Utility function to process files
def process_file(file_path, file_extension):
    try:
        if file_extension == 'pdf':
            return load_pdf(file_path)
        elif file_extension == 'xlsx':
            return load_excel(file_path)
        elif file_extension == 'csv':
            return load_csv(file_path)
        elif file_extension == 'docx':
            return load_docx(file_path)
        else:
            raise ValueError("Unsupported file type")
    except Exception as e:
        logging.error(f"Error processing file {file_path}: {str(e)}")
        return []

# Singleton instance for use across the Flask application
vector_store = VectorStore()




def vectorize_and_search(query, useLLM, k):
    try:
        # Perform similarity search using the singleton vector_store instance
        results = (
            vector_store.chroma_and_LLM_mode(query)
            if useLLM
            else vector_store.pure_chroma_mode(query, k)
        )

        # Accumulate content per file using a single loop
        combined_content = {}
        file_names = []  # List to store filenames for return
        for doc, score in results:
            filename = doc.metadata.get("file_name")
            if filename:  # Ensure filename exists
                file_names.append(filename)
                combined_content[filename] = combined_content.get(filename, "") + doc.page_content + "\n"

        # Generate individual summaries using list comprehension
        summaries = [summarize_data(content) for content in combined_content.values()]

        # Check if only one file is present
        if len(combined_content) == 1:
            full_summary = ""  # No need for a full summary if one file
        else:
            # Generate a full summary from the individual summaries
            full_summary = fullsummarization(summaries)

        return results, file_names, summaries, full_summary

    except Exception as e:
        logging.error(f"Error during vectorization or search: {str(e)}")
        return [], [], [], ""


