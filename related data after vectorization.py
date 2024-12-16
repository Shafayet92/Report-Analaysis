import os
import torch
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain.schema import Document
import numpy as np
from docx import Document as DocxDocument  # Added for handling DOCX files

UPLOAD_FOLDER = './uploads'
PERSIST_DIRECTORY = 'data/db'
SIMILARITY_THRESHOLD = 0.3  # Lowered threshold to accommodate cosine similarity scores
BATCH_SIZE = 500  # Smaller batch size for processing

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def load_pdf(file_path):
    loader = PyPDFLoader(file_path)
    documents = loader.load()
    return documents

def load_excel(file_path):
    df = pd.read_excel(file_path)
    # Process in smaller chunks
    documents = []
    for i in range(0, len(df), BATCH_SIZE):
        chunk = df.iloc[i:i+BATCH_SIZE]
        text = chunk.to_string(index=False)
        documents.append(Document(page_content=text))
    return documents

def load_csv(file_path):
    df = pd.read_csv(file_path)
    # Process in smaller chunks
    documents = []
    for i in range(0, len(df), BATCH_SIZE):
        chunk = df.iloc[i:i+BATCH_SIZE]
        text = chunk.to_string(index=False)
        documents.append(Document(page_content=text))
    return documents

def load_docx(file_path):  # New function to handle DOCX files
    doc = DocxDocument(file_path)
    documents = []
    for para in doc.paragraphs:
        documents.append(Document(page_content=para.text))
    return documents

def text_split(extracted_data):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,  # Reduced chunk size
        chunk_overlap=50,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    text_chunks = text_splitter.split_documents(extracted_data)
    return text_chunks

def get_embeddings():
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={'device': 'cuda' if torch.cuda.is_available() else 'cpu'}
    )

class VectorStore:
    def __init__(self):
        self.embeddings = get_embeddings()
        self.vectordb = None
        self.reset_store()
        self.initialize_store()

    def initialize_store(self):
        if os.path.exists(PERSIST_DIRECTORY):
            self.vectordb = Chroma(
                persist_directory=PERSIST_DIRECTORY,
                embedding_function=self.embeddings
            )
        else:
            self.vectordb = Chroma(
                embedding_function=self.embeddings,
                persist_directory=PERSIST_DIRECTORY
            )

    def add_documents(self, documents):
        # Process documents in batches
        text_chunks = text_split(documents)
        total_chunks = len(text_chunks)

        for i in range(0, total_chunks, BATCH_SIZE):
            batch = text_chunks[i:i + BATCH_SIZE]
            if not self.vectordb:
                self.vectordb = Chroma.from_documents(
                    documents=batch,
                    embedding=self.embeddings,
                    persist_directory=PERSIST_DIRECTORY
                )
            else:
                self.vectordb.add_documents(batch)

            # Chroma should automatically persist to the specified directory without needing explicit save or persist calls.
            print(f"Processed batch {i//BATCH_SIZE + 1}/{(total_chunks + BATCH_SIZE - 1)//BATCH_SIZE}")

        return total_chunks

        # Process documents in batches
        text_chunks = text_split(documents)
        total_chunks = len(text_chunks)

        for i in range(0, total_chunks, BATCH_SIZE):
            batch = text_chunks[i:i + BATCH_SIZE]
            if not self.vectordb:
                self.vectordb = Chroma.from_documents(
                    documents=batch,
                    embedding=self.embeddings,
                    persist_directory=PERSIST_DIRECTORY
                )
            else:
                self.vectordb.add_documents(batch)

            # Instead of calling 'persist', save the vector store to disk.
            self.vectordb.save()  # Save method used for persisting the database.
            print(f"Processed batch {i//BATCH_SIZE + 1}/{(total_chunks + BATCH_SIZE - 1)//BATCH_SIZE}")

        return total_chunks


    def normalize_score(self, score):
        # Convert cosine similarity to [0,1] range
        return (score + 1) / 2

    def similarity_search(self, query, k=100):
        if not self.vectordb:
            return []

        # Handle common misspellings
        query = query.lower().strip()
        common_corrections = {
            'safty': 'safety',
            'quite': 'quit',
            # Add more common corrections as needed
        }
        query = common_corrections.get(query, query)

        results = self.vectordb.similarity_search_with_relevance_scores(query, k=k)

        # Normalize scores and filter
        filtered_results = []
        for doc, score in results:
            normalized_score = self.normalize_score(score)
            if normalized_score >= SIMILARITY_THRESHOLD:
                filtered_results.append((doc, normalized_score))

        return filtered_results

def process_file(file_path, file_extension):
    if file_extension == 'pdf':
        return load_pdf(file_path)
    elif file_extension == 'xlsx':
        return load_excel(file_path)
    elif file_extension == 'csv':
        return load_csv(file_path)
    elif file_extension == 'docx':  # Process DOCX files
        return load_docx(file_path)
    else:
        raise ValueError("Unsupported file type")

def main():
    vector_store = VectorStore()

    # Process files
    for file_name in os.listdir(UPLOAD_FOLDER):
        file_path = os.path.join(UPLOAD_FOLDER, file_name)

        if os.path.isdir(file_path):
            continue

        file_extension = file_name.split('.')[-1].lower()

        try:
            if file_extension in ['pdf', 'xlsx', 'csv', 'docx']:  # Added docx
                print(f"Processing '{file_name}'...")
                documents = process_file(file_path, file_extension)
                chunks_count = vector_store.add_documents(documents)
                print(f"Successfully processed '{file_name}': {chunks_count} chunks vectorized")
            else:
                print(f"Skipping unsupported file type: {file_name}")

        except Exception as e:
            print(f"Error processing '{file_name}': {str(e)}")
            print("Continuing with next file...")

    # Interactive search loop
    while True:
        query = input("\nEnter search query (or 'quit' to exit): ").strip()
        if query.lower() == 'quit':
            break

        results = vector_store.similarity_search(query)

        if not results:
            print("\nNo relevant results found. Try:")
            print("- Using different keywords")
            print("- Checking for spelling mistakes")
            print("- Using more specific terms")
            continue

        print(f"\nFound {len(results)} relevant results:")
        for i, (doc, score) in enumerate(results, 1):
            print(f"\n--- Result {i} (Relevance: {score:.2%}) ---")
            # Clean up the text for display
            content = doc.page_content.replace('_x000D_', '').strip()
            if len(content) > 300:
                content = content[:297] + "..."
            print(content)

if __name__ == '__main__':
    main()
